"""
Fonctions pour l'extraction de texte et d'images, et pour la communication avec ChatGPT
"""
import os
import io
import base64
from openai import OpenAI
from dotenv import load_dotenv
from rag_system import RAGSystem

# Charger les variables d'environnement
load_dotenv()

# Initialiser le client OpenAI
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("OPENAI_API_KEY n'est pas définie dans les variables d'environnement")

client = OpenAI(api_key=api_key)



def describe_image_with_vision(image_base64):
    """
    Décrit une image en utilisant ChatGPT Vision
    
    Args:
        image_base64: Image encodée en base64
        
    Returns:
        str: Description de l'image, ou None en cas d'erreur
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Décris cette image de manière détaillée et précise. Inclus tous les éléments visibles : texte, graphiques, diagrammes, tableaux, etc."
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_base64}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=500
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"[Erreur lors de la description de l'image: {str(e)}]"


def extract_text_from_pdf(file, progress_callback=None):
    """
    Extrait le texte d'un fichier PDF
    
    Args:
        file: Fichier PDF (BytesIO ou file-like object)
        progress_callback: Fonction optionnelle pour afficher la progression (page_num, total_pages, message)
        
    Returns:
        str: Texte extrait du PDF, ou None en cas d'erreur
    """
    try:
        import PyPDF2
        
        pdf_reader = PyPDF2.PdfReader(file)
        total_pages = len(pdf_reader.pages)
        text = ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            if progress_callback:
                progress_callback(page_num + 1, total_pages, f"Traitement de la page {page_num + 1}/{total_pages}...")
            
            # Extraire le texte de la page
            page_text = page.extract_text()
            if page_text.strip():
                text += page_text + "\n"
        
        return text
    
    except ImportError:
        raise ImportError("PyPDF2 n'est pas installé. Installez-le avec: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"Erreur lors de la lecture du PDF: {str(e)}")


def extract_text_from_docx(file):
    """
    Extrait le texte d'un fichier DOCX
    
    Args:
        file: Fichier DOCX (BytesIO ou file-like object)
        
    Returns:
        str: Texte extrait du DOCX, ou None en cas d'erreur
    """
    try:
        from docx import Document
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except ImportError:
        raise ImportError("python-docx n'est pas installé. Installez-le avec: pip install python-docx")
    except Exception as e:
        raise Exception(f"Erreur lors de la lecture du DOCX: {str(e)}")


def extract_text_from_txt(file):
    """
    Extrait le texte d'un fichier TXT
    
    Args:
        file: Fichier TXT (BytesIO ou file-like object)
        
    Returns:
        str: Texte extrait du fichier, ou None en cas d'erreur
    """
    try:
        # Essayer différentes encodages
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                file.seek(0)  # Réinitialiser la position du fichier
                text = file.read().decode(encoding)
                return text
            except UnicodeDecodeError:
                continue
        return file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        raise Exception(f"Erreur lors de la lecture du fichier texte: {str(e)}")


def extract_text(file, file_type, progress_callback=None):
    """
    Extrait le texte selon le type de fichier
    
    Args:
        file: Fichier à traiter (BytesIO ou file-like object)
        file_type: Type MIME du fichier
        progress_callback: Fonction optionnelle pour afficher la progression (pour PDF uniquement)
        
    Returns:
        str: Texte extrait, ou None en cas d'erreur
    """
    if file_type == "application/pdf":
        return extract_text_from_pdf(file, progress_callback)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file)
    elif file_type == "text/plain":
        return extract_text_from_txt(file)
    else:
        raise ValueError(f"Type de fichier non supporté: {file_type}")


def count_pages(file, file_type):
    """
    Compte le nombre de pages d'un document
    
    Args:
        file: Fichier à traiter (BytesIO ou file-like object)
        file_type: Type MIME du fichier
        
    Returns:
        int: Nombre de pages du document
    """
    if file_type == "application/pdf":
        try:
            import PyPDF2
            file.seek(0)  # Réinitialiser la position
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            file.seek(0)  # Réinitialiser pour les prochaines opérations
            return num_pages
        except Exception as e:
            # Si on ne peut pas compter, retourner 0
            return 0
    
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            from docx import Document
            file.seek(0)  # Réinitialiser la position
            doc = Document(file)
            # Estimer le nombre de pages : environ 500 mots par page
            word_count = sum(len(paragraph.text.split()) for paragraph in doc.paragraphs)
            pages = max(1, word_count // 500)  # Au moins 1 page
            return pages
        except Exception:
            return 0
    
    elif file_type == "text/plain":
        try:
            file.seek(0)  # Réinitialiser la position
            text = extract_text_from_txt(file)
            # Estimer : environ 2500 caractères par page
            if text:
                pages = max(1, len(text) // 2500)  # Au moins 1 page
                return pages
            return 0
        except Exception:
            return 0
    
    else:
        return 0



def image_to_base64(image_file):
    """
    Convertit un fichier image en base64
    
    Args:
        image_file: Fichier image (BytesIO ou file-like object)
        
    Returns:
        str: Image encodée en base64, ou None en cas d'erreur
    """
    try:
        image_bytes = image_file.read()
        img_base64 = base64.b64encode(image_bytes).decode('utf-8')
        return img_base64
    except Exception as e:
        raise Exception(f"Erreur lors de la conversion de l'image: {str(e)}")


def get_image_mime_type(image_file):
    """
    Détermine le type MIME d'une image à partir de son nom
    
    Args:
        image_file: Fichier image avec attribut 'name'
        
    Returns:
        str: Type MIME de l'image (image/jpeg ou image/png)
    """
    filename = image_file.name.lower()
    if filename.endswith(('.jpg', '.jpeg')):
        return "image/jpeg"
    elif filename.endswith('.png'):
        return "image/png"
    else:
        return "image/jpeg"  # Par défaut



def ask_question(question, document_text=None, image_base64=None, model="gpt-4o", rag_system=None, use_rag=True):
    """
    Pose une question à ChatGPT avec le contexte du document ou de l'image
    
    Args:
        question: Question de l'utilisateur
        document_text: Texte du document (optionnel)
        image_base64: Image encodée en base64 (optionnel)
        model: Modèle ChatGPT à utiliser
        rag_system: Instance RAGSystem pour la recherche sémantique (optionnel)
        use_rag: Booléen indiquant si RAG doit être utilisé (défaut: True)
        
    Returns:
        str: Réponse de ChatGPT
    """
    try:
        messages = [
            {"role": "system", "content": "Tu es un assistant qui répond aux questions en te basant sur le contenu des documents ou images fournis."}
        ]
        
        # Construire le message utilisateur
        user_content = []
        
        # Si on a une image, utiliser l'API Vision
        if image_base64:
            prompt = f"""Question: {question}

Analyse l'image fournie et répond à la question en te basant uniquement sur le contenu visible dans l'image. Si la réponse n'est pas dans l'image, dis le clairement et dis "je ne sais pas". Réponds dans la même langue que la question de l'utilisateur."""
            
            user_content.append({
                "type": "text",
                "text": prompt
            })
            
            # Le type MIME sera passé séparément, utiliser image/jpeg par défaut
            user_content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{image_base64}"
                }
            })
            
            # Utiliser un modèle avec vision
            vision_models = ["gpt-4o", "gpt-4-turbo", "gpt-4-vision-preview"]
            if model not in vision_models:
                model = "gpt-4o"  # Utiliser gpt-4o par défaut pour la vision
        
        # Si on a du texte de document
        elif document_text:
            # Utiliser RAG si disponible, sinon utiliser tout le texte
            if rag_system and use_rag:
                context = rag_system.get_context_for_question(question, top_k=3)
                
                if context:
                    prompt = f"""Voici les extraits pertinents d'un document trouvés par recherche sémantique :

{context}

Question: {question}

Répond à la question en te basant uniquement sur les extraits du document fournis ci-dessus. Si la réponse n'est pas dans ces extraits, dis le clairement et dis "je ne sais pas". Réponds dans la même langue que la question de l'utilisateur."""
                else:
                    # Fallback si RAG ne trouve rien
                    prompt = f"""Voici le contenu d'un document :

{document_text}

Question: {question}

Répond à la question en te basant uniquement sur le contenu du document fournis. Si la réponse n'est pas dans le document, dis le clairement et dis "je ne sais pas". Réponds dans la même langue que la question de l'utilisateur."""
            else:
                # Mode sans RAG : envoyer tout le document
                prompt = f"""Voici le contenu d'un document :

{document_text}

Question: {question}

Répond à la question en te basant uniquement sur le contenu du document fournis. Si la réponse n'est pas dans le document, dis le clairement et dis "je ne sais pas". Réponds dans la même langue que la question de l'utilisateur."""
            
            user_content.append({"type": "text", "text": prompt})
        
        messages.append({
            "role": "user",
            "content": user_content
        })
        
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=0.2,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Erreur lors de la communication avec ChatGPT: {str(e)}"

